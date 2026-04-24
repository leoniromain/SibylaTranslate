"""
Tradutor de PDF para Word (PT-BR)
Uso: python traduzir_pdf.py <arquivo.pdf> <pagina_inicial> <pagina_final> [saida.docx]
Exemplo: python traduzir_pdf.py meu_livro.pdf 1 5
"""

import sys
import os
import io
import re
import time
import threading
import pdfplumber
import fitz  # PyMuPDF
import numpy as np
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from deep_translator import GoogleTranslator

# ── Configurações ──────────────────────────────────────────────────────────────
SOURCE_LANG  = "en"    # idioma de origem
TARGET_LANG  = "pt"    # português (Brasil)
CHUNK_SIZE   = 4500    # caracteres por chunk (limite Google Translate: 5000)
DELAY_SEC    = 0.5     # pausa entre requisições
IMG_DPI      = 180     # DPI para renderizar páginas-imagem
MAX_IMG_W    = 5.5     # largura máxima da imagem no Word (polegadas)
LINE_GAP_PT  = 8       # gap vertical (pts) que separa parágrafos distintos
OCR_MIN_CONF = 0.35    # confiança mínima do OCR (0-1)

# Palavras/expressões que NUNCA devem ser traduzidas (nomes próprios fixos)
NOMES_PROTEGIDOS: set[str] = set()

# Padrão: sequências de palavras em Title Case (ex: "Paul Huson", "Lord's Prayer")
# com 2+ letras cada — detectados automaticamente
_RE_NOME_PROPRIO = re.compile(
    r'\b([A-Z][a-záàâãéêíóôõúüçñ\'\-]{1,}(?:\s+[A-Z][a-záàâãéêíóôõúüçñ\'\-]{1,})+)\b'
)

# Fontes Windows para desenhar texto traduzido nas imagens
_FONT_PATHS = [
    r"C:\Windows\Fonts\arial.ttf",
    r"C:\Windows\Fonts\calibri.ttf",
    r"C:\Windows\Fonts\times.ttf",
    r"C:\Windows\Fonts\verdana.ttf",
]

# Reader EasyOCR — inicializado uma única vez na primeira chamada
_ocr_reader = None

def get_ocr_reader():
    global _ocr_reader
    if _ocr_reader is None:
        import easyocr
        print("\n  [OCR] Carregando modelo EasyOCR (primeira vez pode demorar)...", end=" ", flush=True)
        _ocr_reader = easyocr.Reader(["en"], gpu=False, verbose=False)
        print("pronto.")
    return _ocr_reader


def _proteger_nomes(texto: str) -> tuple[str, dict]:
    """
    Substitui nomes próprios detectados por placeholders únicos (§0§, §1§…).
    Retorna (texto_com_placeholders, mapa_de_restauração).
    Nomes em NOMES_PROTEGIDOS são sempre protegidos.
    Sequências em Title Case de 2+ palavras são detectadas automaticamente.
    """
    mapa = {}
    idx = [0]  # usa lista para mutabilidade no closure

    def substituir(match):
        nome = match.group(0)
        placeholder = f"\u00a7{idx[0]}\u00a7"
        mapa[placeholder] = nome
        idx[0] += 1
        return placeholder

    # 1. Nomes fixos configurados pelo usuário
    for nome in sorted(NOMES_PROTEGIDOS, key=len, reverse=True):
        escaped = re.escape(nome)
        texto = re.sub(escaped, substituir, texto, flags=re.IGNORECASE)

    # 2. Sequências Title Case automáticas (ex: "Paul Huson", "Lord's Prayer")
    texto = _RE_NOME_PROPRIO.sub(substituir, texto)

    return texto, mapa


def _restaurar_nomes(texto: str, mapa: dict) -> str:
    """Substitui placeholders de volta pelos nomes originais."""
    for placeholder, nome in mapa.items():
        texto = texto.replace(placeholder, nome)
    return texto


def traduzir_texto(texto: str) -> str:
    """Traduz texto preservando nomes próprios e quebras de parágrafo."""
    if not texto.strip():
        return texto

    texto_protegido, mapa = _proteger_nomes(texto)

    translator = GoogleTranslator(source=SOURCE_LANG, target=TARGET_LANG)
    paragrafos = texto_protegido.split("\n")
    resultado = []
    buffer = ""
    for para in paragrafos:
        if len(buffer) + len(para) + 1 > CHUNK_SIZE:
            if buffer.strip():
                try:
                    resultado.extend(translator.translate(buffer).split("\n"))
                    time.sleep(DELAY_SEC)
                except Exception as e:
                    print(f"\n  [AVISO] {e} — mantendo original.")
                    resultado.extend(buffer.split("\n"))
            buffer = para
        else:
            buffer = (buffer + "\n" + para) if buffer else para
    if buffer.strip():
        try:
            resultado.extend(translator.translate(buffer).split("\n"))
            time.sleep(DELAY_SEC)
        except Exception as e:
            print(f"\n  [AVISO] {e} — mantendo original.")
            resultado.extend(buffer.split("\n"))

    traduzido = "\n".join(resultado)
    return _restaurar_nomes(traduzido, mapa)


def pagina_para_imagem_bytes(fitz_doc, page_index: int) -> bytes:
    """Renderiza uma página inteira como PNG e retorna bytes."""
    page = fitz_doc[page_index]
    mat = fitz.Matrix(IMG_DPI / 72, IMG_DPI / 72)
    pix = page.get_pixmap(matrix=mat, alpha=False)
    return pix.tobytes("png")


def para_png_bytes(raw_bytes: bytes, colorspace: str = "") -> bytes:
    """
    Converte qualquer imagem (CMYK, paleta, etc.) para PNG RGB via Pillow.
    Retorna None se não conseguir converter.
    """
    try:
        img = Image.open(io.BytesIO(raw_bytes))
        # CMYK → RGB para compatibilidade com Word
        if img.mode in ("CMYK", "P", "L", "LA"):
            img = img.convert("RGB")
        elif img.mode == "RGBA":
            # Word não lida bem com alpha — compõe sobre branco
            bg = Image.new("RGB", img.size, (255, 255, 255))
            bg.paste(img, mask=img.split()[3])
            img = bg
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        return buf.getvalue()
    except Exception:
        return None


def _carregar_fonte(tamanho: int) -> ImageFont.FreeTypeFont:
    """Tenta carregar uma fonte TTF do sistema. Fallback para fonte padrão."""
    for path in _FONT_PATHS:
        try:
            return ImageFont.truetype(path, max(6, tamanho))
        except Exception:
            pass
    return ImageFont.load_default()


def _cor_contraste(r: int, g: int, b: int) -> tuple:
    """Retorna preto ou branco dependendo da luminância do fundo."""
    lum = 0.299 * r + 0.587 * g + 0.114 * b
    return (0, 0, 0) if lum > 128 else (255, 255, 255)


def _quebrar_texto(texto: str, fonte: ImageFont.FreeTypeFont, largura_max: int) -> list[str]:
    """Quebra o texto em linhas que caibam dentro de largura_max pixels."""
    palavras = texto.split()
    linhas = []
    linha_atual = ""
    for palavra in palavras:
        teste = (linha_atual + " " + palavra).strip()
        bb = fonte.getbbox(teste)
        w = bb[2] - bb[0]
        if w <= largura_max:
            linha_atual = teste
        else:
            if linha_atual:
                linhas.append(linha_atual)
            linha_atual = palavra
    if linha_atual:
        linhas.append(linha_atual)
    return linhas if linhas else [texto]


def _autofit_fonte(texto: str, box_w: int, box_h: int, tam_inicial: int) -> tuple:
    """
    Reduz tamanho da fonte até o texto caber no box com quebra de linha.
    Retorna (fonte, linhas, tamanho, padding).
    """
    PAD = 4
    w_util = max(box_w - PAD * 2, 10)
    h_util = max(box_h - PAD * 2, 6)

    for tamanho in range(min(tam_inicial, 120), 5, -1):
        fonte = _carregar_fonte(tamanho)
        linhas = _quebrar_texto(texto, fonte, w_util)
        line_h = fonte.getbbox("Ay")[3] - fonte.getbbox("Ay")[1]
        altura_total = line_h * len(linhas) + 2 * max(0, len(linhas) - 1)
        if altura_total <= h_util:
            return fonte, linhas, tamanho, PAD

    fonte = _carregar_fonte(6)
    linhas = _quebrar_texto(texto, fonte, w_util)
    return fonte, linhas, 6, PAD


def _agrupar_em_paragrafos(resultados: list, gap_max: int = 12) -> list:
    """
    Agrupa detecções EasyOCR linha a linha em blocos de parágrafo.
    Linhas com gap vertical <= gap_max pixels e x0 similar são do mesmo bloco.
    Retorna lista de grupos: [{texto, x0, y0, x1, y1, conf_media}].
    """
    if not resultados:
        return []

    def bbox_bounds(bbox):
        xs = [int(p[0]) for p in bbox]
        ys = [int(p[1]) for p in bbox]
        return min(xs), min(ys), max(xs), max(ys)

    # Ordena por y_top, depois x0
    itens = []
    for (bbox, texto, conf) in resultados:
        x0, y0, x1, y1 = bbox_bounds(bbox)
        itens.append({"texto": texto, "conf": conf, "x0": x0, "y0": y0, "x1": x1, "y1": y1})
    itens.sort(key=lambda i: (i["y0"], i["x0"]))

    grupos = []
    grupo_atual = [itens[0]]

    for item in itens[1:]:
        ultimo = grupo_atual[-1]
        gap_v  = item["y0"] - ultimo["y1"]
        gap_x  = abs(item["x0"] - grupo_atual[0]["x0"])
        # Mesmo parágrafo: gap vertical pequeno e alinhamento x próximo
        if gap_v <= gap_max and gap_x <= 60:
            grupo_atual.append(item)
        else:
            grupos.append(grupo_atual)
            grupo_atual = [item]

    grupos.append(grupo_atual)

    resultado_grupos = []
    for g in grupos:
        texto_completo = " ".join(i["texto"] for i in g)
        conf_media = sum(i["conf"] for i in g) / len(g)
        gx0 = min(i["x0"] for i in g)
        gy0 = min(i["y0"] for i in g)
        gx1 = max(i["x1"] for i in g)
        gy1 = max(i["y1"] for i in g)
        resultado_grupos.append({
            "texto": texto_completo,
            "conf":  conf_media,
            "x0": gx0, "y0": gy0, "x1": gx1, "y1": gy1,
        })

    return resultado_grupos


def ocr_traduzir_imagem(img_bytes: bytes) -> bytes:
    """
    Faz OCR na imagem, agrupa linhas em parágrafos, traduz cada bloco e
    redesenha o texto traduzido com auto-fit de fonte e quebra de linha.
    Retorna PNG bytes com o texto já em PT-BR.
    """
    reader = get_ocr_reader()
    img = Image.open(io.BytesIO(img_bytes)).convert("RGB")
    img_np = np.array(img)

    resultados_raw = reader.readtext(img_np, paragraph=False)
    if not resultados_raw:
        return img_bytes

    # Filtra baixa confiança antes de agrupar
    resultados_raw = [(b, t, c) for (b, t, c) in resultados_raw
                      if c >= OCR_MIN_CONF and t.strip()]
    if not resultados_raw:
        return img_bytes

    grupos = _agrupar_em_paragrafos(resultados_raw)
    draw = ImageDraw.Draw(img)

    for g in grupos:
        x0, y0, x1, y1 = g["x0"], g["y0"], g["x1"], g["y1"]
        box_w = max(x1 - x0, 30)
        box_h = max(y1 - y0, 12)

        # Amostra a cor de fundo ANTES de apagar (mediana = mais robusta)
        regiao_arr = np.array(img.crop((x0, y0, x1, y1)))
        cor_fundo  = tuple(int(np.median(regiao_arr[:, :, c])) for c in range(3))
        cor_texto  = _cor_contraste(*cor_fundo)

        # Traduz o parágrafo completo (uma única chamada ao tradutor)
        traduzido = traduzir_texto(g["texto"])

        # Tamanho inicial baseado na altura de uma linha do grupo original
        n_linhas_orig = max(1, len(g["texto"].split("\n")))
        tam_inicial = max((box_h // n_linhas_orig) - 4, 8)

        fonte, linhas, _, pad = _autofit_fonte(traduzido, box_w, box_h, tam_inicial)

        # Apaga toda a área do parágrafo
        draw.rectangle([x0, y0, x1, y1], fill=cor_fundo)

        # Renderiza centralizado verticalmente
        line_h = fonte.getbbox("Ay")[3] - fonte.getbbox("Ay")[1]
        total_h = line_h * len(linhas) + 2 * max(0, len(linhas) - 1)
        y_cur = y0 + max(pad, (box_h - total_h) // 2)

        for linha in linhas:
            draw.text((x0 + pad, y_cur), linha, fill=cor_texto, font=fonte)
            y_cur += line_h + 2

    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def extrair_imagens_da_pagina(fitz_doc, page_index: int):
    """
    Retorna lista de dicts {bytes, y_top} com imagens embutidas na página,
    ordenadas por posição vertical. Converte tudo para PNG.
    """
    page = fitz_doc[page_index]
    imagens = []
    for img_info in page.get_images(full=True):
        xref = img_info[0]
        # Posição da imagem na página via get_image_bbox
        try:
            bbox = page.get_image_bbox(img_info)
            y_top = bbox.y0
        except Exception:
            y_top = 0
        try:
            base_img = fitz_doc.extract_image(xref)
            raw = base_img["image"]
            cs  = base_img.get("colorspace", "")
            png = para_png_bytes(raw, cs)
            if png:
                png = ocr_traduzir_imagem(png)  # OCR + tradução na imagem
                imagens.append({"bytes": png, "y_top": y_top})
        except Exception:
            pass
    imagens.sort(key=lambda x: x["y_top"])
    return imagens


def extrair_blocos_pagina(page):
    """
    Extrai blocos de texto agrupando linhas próximas em parágrafos.
    Cada bloco: {texto, tamanho, negrito, italico, x0, y_top, y_bot}
    """
    words = page.extract_words(
        x_tolerance=3,
        y_tolerance=3,
        keep_blank_chars=False,
        use_text_flow=True,
        extra_attrs=["size", "fontname"],
    )
    if not words:
        return []

    # Agrupa palavras em linhas pelo y
    linhas_dict = {}
    for w in words:
        y_key = round(w["top"] / 3) * 3
        linhas_dict.setdefault(y_key, []).append(w)

    # Constrói lista de linhas ordenadas
    linhas = []
    for y_key in sorted(linhas_dict.keys()):
        ws = sorted(linhas_dict[y_key], key=lambda w: w["x0"])
        texto = " ".join(w["text"] for w in ws)
        primeiro = ws[0]
        tamanho  = primeiro.get("size", 11) or 11
        fontname = (primeiro.get("fontname") or "").lower()
        linhas.append({
            "texto":   texto,
            "tamanho": tamanho,
            "negrito": "bold" in fontname or "heavy" in fontname,
            "italico": "italic" in fontname or "oblique" in fontname,
            "x0":      primeiro["x0"],
            "y_top":   primeiro["top"],
            "y_bot":   ws[-1].get("bottom", primeiro["top"] + tamanho),
        })

    # Agrupa linhas consecutivas em parágrafos (gap pequeno = mesmo parágrafo)
    if not linhas:
        return []

    blocos = []
    atual = dict(linhas[0])

    for prox in linhas[1:]:
        gap = prox["y_top"] - atual["y_bot"]
        mesmo_estilo = (
            abs(prox["tamanho"] - atual["tamanho"]) < 1.5
            and prox["negrito"] == atual["negrito"]
            and prox["italico"] == atual["italico"]
            and abs(prox["x0"] - atual["x0"]) < 30
        )
        if gap <= LINE_GAP_PT and mesmo_estilo:
            # Mesmo parágrafo: concatena
            atual["texto"] += " " + prox["texto"]
            atual["y_bot"] = prox["y_bot"]
        else:
            blocos.append(atual)
            atual = dict(prox)

    blocos.append(atual)
    return blocos


def inserir_imagem_no_doc(doc, img_bytes: bytes, centralizar: bool = True):
    """Insere uma imagem PNG (bytes) no documento Word."""
    png = para_png_bytes(img_bytes)  # garante PNG mesmo se já renderizado
    if not png:
        return
    para = doc.add_paragraph()
    if centralizar:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(io.BytesIO(png), width=Inches(MAX_IMG_W))


def adicionar_pagina_no_doc(doc, num_pagina, blocos_traduzidos, imagens_embutidas):
    """
    Intercala blocos de texto traduzidos e imagens embutidas
    no documento Word, respeitando a ordem vertical.
    """
    # Cabeçalho de página
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"— Página {num_pagina} —")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True

    # Intercala texto e imagens por posição vertical
    itens = []
    for bloco in blocos_traduzidos:
        itens.append(("texto", bloco["y_top"], bloco))
    for img in imagens_embutidas:
        itens.append(("imagem", img["y_top"], img))
    itens.sort(key=lambda x: x[1])

    for tipo, _, dado in itens:
        if tipo == "imagem":
            try:
                inserir_imagem_no_doc(doc, dado["bytes"])
            except Exception as e:
                p = doc.add_paragraph(f"[Imagem não pôde ser inserida: {e}]")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            texto = dado["texto"].strip()
            if not texto:
                doc.add_paragraph()
                continue

            para = doc.add_paragraph()
            font_size = max(8, min(round(dado["tamanho"]), 36))

            # Indentação proporcional
            if dado["x0"] > 60:
                para.paragraph_format.left_indent = Pt(dado["x0"] * 0.4)

            # Heurística de título
            if font_size >= 14 and len(texto) < 80:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            run = para.add_run(texto)
            run.font.size = Pt(font_size)
            run.bold   = dado["negrito"]
            run.italic = dado["italico"]

    doc.add_page_break()


def _encontrar_marcadores(doc) -> dict:
    """
    Varre os parágrafos do doc e retorna:
      {num_pagina: (idx_inicio, idx_fim_exclusivo)}
    onde idx_inicio é o índice do parágrafo marcador "— Página N —"
    e idx_fim é o início do próximo marcador (ou total de parágrafos).
    """
    padrao = re.compile(r"— Página (\d+) —")
    marcadores = {}  # num_pagina -> idx_paragrafo
    for i, para in enumerate(doc.paragraphs):
        m = padrao.search(para.text)
        if m:
            marcadores[int(m.group(1))] = i

    nums = sorted(marcadores)
    total = len(doc.paragraphs)
    ranges = {}
    for j, num in enumerate(nums):
        ini = marcadores[num]
        fim = marcadores[nums[j + 1]] if j + 1 < len(nums) else total
        ranges[num] = (ini, fim)
    return ranges


def _remover_pagina_do_doc(doc, num_pagina) -> object:
    """
    Remove todos os parágrafos da seção "Página N" do doc.
    Retorna o elemento XML de referência ANTES do qual reinserir conteúdo,
    ou None se a página não foi encontrada (append ao final).
    """
    ranges = _encontrar_marcadores(doc)
    if num_pagina not in ranges:
        return None

    ini, fim = ranges[num_pagina]
    paras = doc.paragraphs

    # Elemento de referência = primeiro parágrafo da página SEGUINTE
    ref_element = paras[fim]._element if fim < len(paras) else None

    # Remove de trás para frente para não deslocar índices
    for i in range(fim - 1, ini - 1, -1):
        elem = paras[i]._element
        elem.getparent().remove(elem)

    return ref_element  # pode ser None se era a última página


def _mover_para_antes_de(doc, count_before: int, ref_element):
    """
    Move os parágrafos adicionados após count_before para antes de ref_element.
    Usado para inserir conteúdo de replace na posição correta.
    """
    if ref_element is None:
        return  # já está no final — sem necessidade de mover
    while len(doc.paragraphs) > count_before:
        elem = doc.paragraphs[count_before]._element
        ref_element.addprevious(elem)


def _processar_pagina(num, pdf, fitz_doc, doc, modo, ref_element=None):
    """
    Extrai, traduz e insere uma página no doc.
    No modo 'replace', move o conteúdo inserido para a posição correta.
    """
    page   = pdf.pages[num - 1]
    blocos = extrair_blocos_pagina(page)
    imagens = extrair_imagens_da_pagina(fitz_doc, num - 1)

    count_before = len(doc.paragraphs)

    if not blocos and not imagens:
        print("sem texto/imagens → renderizando + OCR...", end=" ", flush=True)
        img_bytes = pagina_para_imagem_bytes(fitz_doc, num - 1)
        img_bytes = ocr_traduzir_imagem(img_bytes)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"— Página {num} —")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run.italic = True
        inserir_imagem_no_doc(doc, img_bytes)
        doc.add_page_break()
    else:
        print(f"{len(blocos)} blocos, {len(imagens)} imagens. Traduzindo...", end=" ", flush=True)
        blocos_traduzidos = []
        for bloco in blocos:
            blocos_traduzidos.append({**bloco, "texto": traduzir_texto(bloco["texto"])})
        adicionar_pagina_no_doc(doc, num, blocos_traduzidos, imagens)

    if modo == "replace" and ref_element is not None:
        _mover_para_antes_de(doc, count_before, ref_element)

    print("OK")


def processar(pdf_path: str, pag_ini: int, pag_fim: int, saida: str,
              modo: str = "novo", arquivo_base: str = None,
              cancel_event: threading.Event | None = None):
    """
    Modos:
      novo    — cria novo arquivo .docx
      append  — abre arquivo existente e adiciona páginas ao final
      replace — abre arquivo existente e substitui as páginas indicadas no lugar certo
    """
    print(f"\nArquivo : {pdf_path}")
    print(f"Páginas : {pag_ini} a {pag_fim}")
    print(f"Modo    : {modo.upper()}")
    print(f"Saída   : {saida}\n")

    if modo in ("append", "replace"):
        if not arquivo_base or not os.path.isfile(arquivo_base):
            print(f"Erro: arquivo base '{arquivo_base}' não encontrado.")
            sys.exit(1)
        doc = Document(arquivo_base)
        print(f"Abrindo base: {arquivo_base}")
        if modo == "replace":
            ranges = _encontrar_marcadores(doc)
            paginas_existentes = sorted(ranges.keys())
            print(f"Páginas encontradas no doc: {paginas_existentes[0]}–{paginas_existentes[-1]}" if paginas_existentes else "  (nenhuma página marcada)")
    else:
        doc = Document()
        for section in doc.sections:
            section.top_margin    = Pt(72)
            section.bottom_margin = Pt(72)
            section.left_margin   = Pt(72)
            section.right_margin  = Pt(72)

    fitz_doc = fitz.open(pdf_path)

    with pdfplumber.open(pdf_path) as pdf:
        total   = len(pdf.pages)
        pag_fim = min(pag_fim, total)
        print(f"Total de páginas no PDF: {total}\n")

        for num in range(pag_ini, pag_fim + 1):
            if cancel_event is not None and cancel_event.is_set():
                print("\n⚠️  Tradução cancelada pelo usuário.")
                break
            print(f"[{num}/{pag_fim}] Página {num}...", end=" ", flush=True)

            if modo == "replace":
                ref_element = _remover_pagina_do_doc(doc, num)
                if ref_element is None and num in _encontrar_marcadores(doc):
                    # era a última página — ref_element None significa "ao final"
                    pass
                _processar_pagina(num, pdf, fitz_doc, doc, modo, ref_element)
            else:
                _processar_pagina(num, pdf, fitz_doc, doc, modo)

    fitz_doc.close()
    doc.save(saida)
    print(f"\nConcluído! Arquivo salvo em: {saida}")


# ── Entry point ────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(
        description="Traduz páginas de um PDF para PT-BR e salva em Word.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Exemplos:
  # Traduz páginas 1-5 → novo arquivo
  python traduzir_pdf.py livro.pdf 1 5

  # Traduz páginas 1-5 com nome personalizado
  python traduzir_pdf.py livro.pdf 1 5 -o minha_traducao.docx

  # Continua tradução (páginas 6-10) no arquivo existente
  python traduzir_pdf.py livro.pdf 6 10 --append minha_traducao.docx

  # Re-traduz a página 5 e a substitui no arquivo existente
  python traduzir_pdf.py livro.pdf 5 5 --replace minha_traducao.docx
        """
    )
    parser.add_argument("pdf",      help="Caminho para o arquivo PDF")
    parser.add_argument("pag_ini",  type=int, help="Página inicial")
    parser.add_argument("pag_fim",  type=int, help="Página final")
    parser.add_argument("-o", "--output", help="Nome do arquivo de saída .docx")

    grupo = parser.add_mutually_exclusive_group()
    grupo.add_argument("--append",  metavar="DOCX",
                       help="Continua adicionando páginas ao final do arquivo DOCX informado")
    grupo.add_argument("--replace", metavar="DOCX",
                       help="Substitui as páginas no arquivo DOCX informado (mantém posição)")

    args = parser.parse_args()

    if not os.path.isfile(args.pdf):
        print(f"Erro: arquivo '{args.pdf}' não encontrado.")
        sys.exit(1)

    if args.pag_ini < 1 or args.pag_fim < args.pag_ini:
        print("Erro: intervalo de páginas inválido.")
        sys.exit(1)

    # Determina modo e arquivo base
    if args.append:
        modo         = "append"
        arquivo_base = args.append
        saida        = args.output or args.append
    elif args.replace:
        modo         = "replace"
        arquivo_base = args.replace
        saida        = args.output or args.replace
    else:
        modo         = "novo"
        arquivo_base = None
        base         = os.path.splitext(os.path.basename(args.pdf))[0]
        saida        = args.output or f"{base}_pt_p{args.pag_ini}-p{args.pag_fim}.docx"

    processar(args.pdf, args.pag_ini, args.pag_fim, saida, modo, arquivo_base)

