import io
import re
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .image_utils import para_png_bytes, MAX_IMG_W


def inserir_imagem_no_doc(doc, img_bytes: bytes, centralizar: bool = True) -> None:
    """Insere uma imagem PNG (bytes) no documento Word."""
    png = para_png_bytes(img_bytes)
    if not png:
        return
    para = doc.add_paragraph()
    if centralizar:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run().add_picture(io.BytesIO(png), width=Inches(MAX_IMG_W))


def adicionar_pagina_no_doc(doc, num_pagina: int, blocos_traduzidos: list,
                             imagens_embutidas: list) -> None:
    """
    Intercala blocos de texto traduzidos e imagens embutidas no documento Word,
    respeitando a ordem vertical.
    """
    # Cabeçalho de página
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"— Página {num_pagina} —")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True

    itens = (
        [("texto",  bloco["y_top"], bloco) for bloco in blocos_traduzidos] +
        [("imagem", img["y_top"],   img)   for img   in imagens_embutidas]
    )
    itens.sort(key=lambda x: x[1])

    for tipo, _, dado in itens:
        if tipo == "imagem":
            try:
                inserir_imagem_no_doc(doc, dado["bytes"])
            except Exception as e:
                p = doc.add_paragraph(f"[Imagem não pôde ser inserida: {e}]")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            texto = dado["texto"].strip()
            if not texto:
                doc.add_paragraph()
                continue
            para = doc.add_paragraph()
            font_size = max(8, min(round(dado["tamanho"]), 36))
            if dado["x0"] > 60:
                para.paragraph_format.left_indent = Pt(dado["x0"] * 0.4)
            if font_size >= 14 and len(texto) < 80:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(texto)
            run.font.size = Pt(font_size)
            run.bold   = dado["negrito"]
            run.italic = dado["italico"]

    doc.add_page_break()


def _encontrar_marcadores(doc) -> dict:
    """
    Varre os parágrafos e retorna {num_pagina: (idx_inicio, idx_fim_exclusivo)}.
    """
    padrao = re.compile(r"— Página (\d+) —")
    marcadores: dict = {}
    for i, para in enumerate(doc.paragraphs):
        m = padrao.search(para.text)
        if m:
            marcadores[int(m.group(1))] = i

    nums = sorted(marcadores)
    total = len(doc.paragraphs)
    ranges: dict = {}
    for j, num in enumerate(nums):
        ini = marcadores[num]
        fim = marcadores[nums[j + 1]] if j + 1 < len(nums) else total
        ranges[num] = (ini, fim)
    return ranges


def _remover_pagina_do_doc(doc, num_pagina: int) -> object:
    """
    Remove todos os parágrafos da seção "Página N".
    Retorna o elemento XML de referência antes do qual reinserir,
    ou None se a página não foi encontrada.
    """
    ranges = _encontrar_marcadores(doc)
    if num_pagina not in ranges:
        return None

    ini, fim = ranges[num_pagina]
    paras = doc.paragraphs
    ref_element = paras[fim]._element if fim < len(paras) else None

    for i in range(fim - 1, ini - 1, -1):
        paras[i]._element.getparent().remove(paras[i]._element)

    return ref_element


def _mover_para_antes_de(doc, count_before: int, ref_element) -> None:
    """Move os parágrafos adicionados após count_before para antes de ref_element."""
    if ref_element is None:
        return
    while len(doc.paragraphs) > count_before:
        elem = doc.paragraphs[count_before]._element
        ref_element.addprevious(elem)
