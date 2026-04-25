import sys
import os
import threading

import pdfplumber
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .translation import traduzir_texto
from .image_utils import pagina_para_imagem_bytes
from .ocr import ocr_traduzir_imagem
from .pdf_reader import extrair_blocos_pagina, extrair_imagens_da_pagina
from .word_writer import (
    inserir_imagem_no_doc,
    adicionar_pagina_no_doc,
    _encontrar_marcadores,
    _remover_pagina_do_doc,
    _mover_para_antes_de,
)
from .text_writer import salvar_txt, salvar_md
from .pdf_writer import salvar_pdf

_FORMATOS_TEXTO = {"txt", "md", "pdf"}  # formatos que não usam python-docx


def _extrair_traduzir_pagina(num: int, pdf, fitz_doc,
                             lang_src: str, lang_dst: str) -> dict:
    """
    Extrai e traduz uma página. Retorna dict com blocos e imagens já traduzidos.
    Usado pelos formatos texto (txt/md) onde não há objeto doc.
    """
    page    = pdf.pages[num - 1]
    blocos  = extrair_blocos_pagina(page)
    imagens = extrair_imagens_da_pagina(fitz_doc, num - 1, lang_src, lang_dst)

    if not blocos and not imagens:
        print("sem texto/imagens → renderizando + OCR...", end=" ", flush=True)
        img_bytes = pagina_para_imagem_bytes(fitz_doc, num - 1)
        img_bytes = ocr_traduzir_imagem(img_bytes, lang_src, lang_dst)
        blocos_traduzidos: list = []
        imagens_ret = [{"bytes": img_bytes, "y_top": 0}]
    else:
        print(f"{len(blocos)} blocos, {len(imagens)} imagens. Traduzindo...",
              end=" ", flush=True)
        blocos_traduzidos = [
            {**b, "texto": traduzir_texto(b["texto"], lang_src, lang_dst)}
            for b in blocos
        ]
        imagens_ret = imagens

    print("OK")
    return {"num": num, "blocos": blocos_traduzidos, "imagens": imagens_ret}


def _processar_pagina(num: int, pdf, fitz_doc, doc, modo: str,
                      ref_element=None,
                      lang_src: str = "en", lang_dst: str = "pt") -> None:
    """Extrai, traduz e insere uma página no doc."""
    page    = pdf.pages[num - 1]
    blocos  = extrair_blocos_pagina(page)
    imagens = extrair_imagens_da_pagina(fitz_doc, num - 1, lang_src, lang_dst)
    count_before = len(doc.paragraphs)

    if not blocos and not imagens:
        print("sem texto/imagens → renderizando + OCR...", end=" ", flush=True)
        img_bytes = pagina_para_imagem_bytes(fitz_doc, num - 1)
        img_bytes = ocr_traduzir_imagem(img_bytes, lang_src, lang_dst)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"— Página {num} —")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run.italic = True
        inserir_imagem_no_doc(doc, img_bytes)
        doc.add_page_break()
    else:
        print(f"{len(blocos)} blocos, {len(imagens)} imagens. Traduzindo...",
              end=" ", flush=True)
        blocos_traduzidos = [
            {**bloco, "texto": traduzir_texto(bloco["texto"], lang_src, lang_dst)}
            for bloco in blocos
        ]
        adicionar_pagina_no_doc(doc, num, blocos_traduzidos, imagens)

    if modo == "replace" and ref_element is not None:
        _mover_para_antes_de(doc, count_before, ref_element)

    print("OK")


def processar(pdf_path: str, pag_ini: int, pag_fim: int, saida: str,
              modo: str = "novo", arquivo_base: str | None = None,
              cancel_event: threading.Event | None = None,
              lang_src: str = "en", lang_dst: str = "pt",
              fmt: str = "docx") -> None:
    """
    Processa um PDF e salva tradução no formato escolhido.

    fmt: 'docx' (padrão) | 'txt' | 'md'
    Modos (só aplicáveis para fmt='docx'):
      novo    — cria novo arquivo .docx
      append  — abre arquivo existente e adiciona páginas ao final
      replace — abre arquivo existente e substitui as páginas indicadas
    """
    print(f"\nArquivo : {pdf_path}")
    print(f"Páginas : {pag_ini} a {pag_fim}")
    print(f"Modo    : {modo.upper()}")
    print(f"Idiomas : {lang_src} → {lang_dst}")
    print(f"Formato : {fmt.upper()}")
    print(f"Saída   : {saida}\n")

    # ── Formatos texto puro (txt / md) ────────────────────────────────────────
    if fmt in _FORMATOS_TEXTO:
        fitz_doc = fitz.open(pdf_path)
        paginas_data: list[dict] = []
        with pdfplumber.open(pdf_path) as pdf:
            total   = len(pdf.pages)
            pag_fim = min(pag_fim, total)
            print(f"Total de páginas no PDF: {total}\n")
            for num in range(pag_ini, pag_fim + 1):
                if cancel_event is not None and cancel_event.is_set():
                    print("\n⚠️  Tradução cancelada pelo usuário.")
                    break
                print(f"[{num}/{pag_fim}] Página {num}...", end=" ", flush=True)
                paginas_data.append(
                    _extrair_traduzir_pagina(num, pdf, fitz_doc, lang_src, lang_dst)
                )
        fitz_doc.close()
        if fmt == "txt":
            salvar_txt(paginas_data, saida)
        elif fmt == "md":
            salvar_md(paginas_data, saida)
        else:  # pdf
            salvar_pdf(paginas_data, saida)
        print(f"\nConcluído! Arquivo salvo em: {saida}")
        return

    # ── Formato Word (.docx) ──────────────────────────────────────────────────
    if modo in ("append", "replace"):
        if not arquivo_base or not os.path.isfile(arquivo_base):
            print(f"Erro: arquivo base '{arquivo_base}' não encontrado.")
            sys.exit(1)
        doc = Document(arquivo_base)
        print(f"Abrindo base: {arquivo_base}")
        if modo == "replace":
            ranges = _encontrar_marcadores(doc)
            paginas_existentes = sorted(ranges.keys())
            print(
                f"Páginas encontradas no doc: "
                f"{paginas_existentes[0]}–{paginas_existentes[-1]}"
                if paginas_existentes else "  (nenhuma página marcada)"
            )
    else:
        doc = Document()
        for section in doc.sections:
            section.top_margin    = Pt(72)
            section.bottom_margin = Pt(72)
            section.left_margin   = Pt(72)
            section.right_margin  = Pt(72)

    fitz_doc = fitz.open(pdf_path)

    with pdfplumber.open(pdf_path) as pdf:
        total   = len(pdf.pages)
        pag_fim = min(pag_fim, total)
        print(f"Total de páginas no PDF: {total}\n")

        for num in range(pag_ini, pag_fim + 1):
            if cancel_event is not None and cancel_event.is_set():
                print("\n⚠️  Tradução cancelada pelo usuário.")
                break
            print(f"[{num}/{pag_fim}] Página {num}...", end=" ", flush=True)

            if modo == "replace":
                ref_element = _remover_pagina_do_doc(doc, num)
                _processar_pagina(num, pdf, fitz_doc, doc, modo, ref_element,
                                  lang_src, lang_dst)
            else:
                _processar_pagina(num, pdf, fitz_doc, doc, modo,
                                  lang_src=lang_src, lang_dst=lang_dst)

    fitz_doc.close()
    doc.save(saida)
    print(f"\nConcluído! Arquivo salvo em: {saida}")
